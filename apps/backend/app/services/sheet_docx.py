"""Word (.docx) sheet export (X4, AD-26).

python-docx builds the document. Math ``question_latex`` is rendered to a PNG
via matplotlib mathtext and embedded (Word cannot consume raw LaTeX); English
``question_text`` goes in as a plain paragraph. Questions whose transcribed text
is missing fall back to the cropped question screenshot. Sheets never include
answer keys — ``solution_note`` / ``error_category`` are deliberately omitted.
"""

import logging
import os
import tempfile
import uuid

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.sheet_text import (
    ANSWER_HINT,
    SUBJECT_LABELS,
    TITLE,
    TYPE_LABELS,
    SheetQuestionData,
)

logger = logging.getLogger(__name__)

DOCX_DIR = "data/images/sheets"
DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

_TITLE_COLOR = RGBColor(0x1E, 0x1B, 0x18)  # #1E1B18
_SUBTITLE_COLOR = RGBColor(0x6B, 0x65, 0x60)  # #6B6560
_CJK_FONT = "宋体"


def render_latex_png(latex: str, output_path: str, dpi: int = 150) -> str:
    """Render a LaTeX math expression to a PNG using matplotlib mathtext.

    Raises ``ValueError`` (or other ``Exception``) on unsupported markup so the
    caller can fall back to the question screenshot. Matplotlib is imported
    lazily: it is a heavy dependency used only for math-question sheets.
    """
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    expr = _strip_delimiters(latex)
    if not expr:
        raise ValueError("empty LaTeX expression")

    fig = plt.figure()
    fig.text(0, 0, f"${expr}$", fontsize=14)
    fig.savefig(
        output_path,
        dpi=dpi,
        bbox_inches="tight",
        pad_inches=0.05,
        transparent=True,
    )
    plt.close(fig)
    return output_path


def _strip_delimiters(latex: str) -> str:
    expr = latex.strip()
    for open_delim, close_delim in (("$$", "$$"), ("\\[", "\\]"), ("$", "$")):
        if expr.startswith(open_delim) and expr.endswith(close_delim):
            expr = expr[len(open_delim) : len(expr) - len(close_delim)]
            break
    return expr.strip()


def _set_cjk_font(run, size: int, bold: bool = False, color=None) -> None:
    """Set an East-Asian font so Chinese text renders correctly in Word."""
    run.font.name = _CJK_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _CJK_FONT)


def _add_answer_box(doc: Document, height_cm: float = 4.0) -> None:
    """Append a bordered, fixed-height answer area below a question."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)

    # Hint paragraph inside the box (grey, unobtrusive).
    cell.paragraphs[0].text = ""
    hint = cell.paragraphs[0].add_run(f"（{ANSWER_HINT}）")
    _set_cjk_font(hint, 9, color=_SUBTITLE_COLOR)

    # Enforce a minimum row height.
    tr = table.rows[0]._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(Cm(height_cm).twips)))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def _add_fallback_image(doc: Document, question: SheetQuestionData) -> None:
    """Embed the cropped question screenshot for a text-less question."""
    path = question.question_image_path
    if path and os.path.isfile(path):
        doc.add_picture(path, width=Cm(10))
    else:
        para = doc.add_paragraph("（题干缺失，请参考原题截图）")
        _set_cjk_font(para.runs[0], 10, color=_SUBTITLE_COLOR)


def _add_question_body(
    doc: Document, question: SheetQuestionData, tmp_dir: str, index: int
) -> None:
    """Render a question's body: text, LaTeX PNG, or fallback screenshot."""
    if question.is_incomplete:
        _add_fallback_image(doc, question)
        return

    if question.subject == "math":
        png_path = os.path.join(tmp_dir, f"q_{index}.png")
        try:
            render_latex_png(question.question_latex or "", png_path)
            doc.add_picture(png_path, width=Cm(8))
            return
        except Exception as exc:  # noqa: BLE001 — mathtext is best-effort
            logger.warning(
                "LaTeX render failed for question %s, falling back to screenshot: %s",
                question.question_number,
                exc,
            )

    else:
        para = doc.add_paragraph(question.question_text or "")
        _set_cjk_font(para.runs[0], 12, color=_TITLE_COLOR)
        return

    # Math render failure with no usable text → screenshot.
    _add_fallback_image(doc, question)


def build_sheet_docx(
    questions: list[SheetQuestionData],
    child_name: str,
    subject: str,
    output_dir: str = DOCX_DIR,
) -> str:
    """Build a .docx text sheet and return the path it was written to.

    ``questions`` is the ordered list of selected entries. The title bar carries
    the child name + subject + date; each question is followed by a bordered
    answer area. No answer key is included.
    """
    from datetime import datetime, timezone

    subject_label = SUBJECT_LABELS.get(subject, subject)
    today = datetime.now(timezone.utc).strftime("%Y年%m月%d日")

    doc = Document()

    # Title bar.
    title = doc.add_heading(level=0)
    run = title.add_run(TITLE)
    _set_cjk_font(run, 20, bold=True, color=_TITLE_COLOR)
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(f"{child_name}  |  {subject_label}  |  {today}")
    _set_cjk_font(sub_run, 12, color=_SUBTITLE_COLOR)

    # Question sections.
    with tempfile.TemporaryDirectory() as tmp_dir:
        for index, question in enumerate(questions, start=1):
            type_label = TYPE_LABELS.get(question.question_type, question.question_type)
            label = doc.add_paragraph()
            label_run = label.add_run(
                f"第 {question.question_number} 题  [{type_label}]"
            )
            _set_cjk_font(label_run, 12, bold=True, color=_TITLE_COLOR)

            _add_question_body(doc, question, tmp_dir, index)
            _add_answer_box(doc)

        os.makedirs(output_dir, exist_ok=True)
        filename = f"{uuid.uuid4().hex}.docx"
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)

    logger.info(
        "Practice sheet docx saved: %s (%d questions)", output_path, len(questions)
    )
    return output_path
