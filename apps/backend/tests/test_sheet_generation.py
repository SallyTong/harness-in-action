"""Tests for X4: text sheet assembly + .docx export + generate `format` switch."""

import os
import time

import pytest
from docx import Document
from PIL import Image

from app.models.error_question import ErrorQuestion
from app.models.submission import Submission
from app.routers import error_collections as ec_router
from app.services.image_signing import sign
from app.services.sheet_docx import build_sheet_docx, render_latex_png
from app.services.sheet_text import SheetQuestionData, assemble_sheet_questions
from tests.helpers import login

PHONE = "13800138000"

# LaTeX that matplotlib mathtext can render without external fonts/TeX.
LATEX_FRACTION = r"\frac{1}{2} + \frac{3}{4}"


def _auth(token: str) -> dict[str, str]:
    return {"Authorization": f"Bearer {token}"}


async def _get_child_id(client, token: str) -> int:
    resp = await client.get("/api/children", headers=_auth(token))
    assert resp.status_code == 200
    return resp.json()[0]["id"]


async def _seed_error_questions(db, child_id: int, subject: str, specs) -> int:
    """Create a Submission + N ErrorQuestion rows; return submission id."""
    sub = Submission(
        child_id=child_id,
        subject=subject,
        status="completed",
        original_image_path="data/images/originals/test.jpg",
    )
    db.add(sub)
    await db.flush()
    for i, spec in enumerate(specs, start=1):
        db.add(
            ErrorQuestion(
                submission_id=sub.id,
                child_id=child_id,
                subject=subject,
                question_number=spec.get("number", str(i)),
                question_type=spec.get("type", "calculation"),
                question_image_path=spec.get(
                    "image", f"data/images/questions/{sub.id}_{i}.jpg"
                ),
                question_text=spec.get("text"),
                question_latex=spec.get("latex"),
            )
        )
    await db.flush()
    return sub.id


# ── Assembly (sheet_text) ─────────────────────────────────────────────


def test_assemble_questions_resolves_math_latex_and_flags_incomplete():
    raw = [
        _fake_eq(number="1", subject="math", latex=LATEX_FRACTION),
        _fake_eq(number="2", subject="math", latex=None),  # incomplete → fallback
    ]
    entries = assemble_sheet_questions(raw)
    assert len(entries) == 2
    assert entries[0].primary_text == LATEX_FRACTION
    assert entries[0].is_incomplete is False
    assert entries[1].is_incomplete is True


def test_assemble_questions_english_uses_text_not_latex():
    entries = assemble_sheet_questions(
        [_fake_eq(number="1", subject="english", text="What is your name?")]
    )
    assert entries[0].primary_text == "What is your name?"
    assert entries[0].is_incomplete is False


class _fake_eq:
    """Minimal stand-in for an ErrorQuestion ORM object."""

    def __init__(self, *, number, subject, text=None, latex=None, image="q.jpg"):
        self.question_number = number
        self.subject = subject
        self.question_text = text
        self.question_latex = latex
        self.question_image_path = image
        self.question_type = "calculation"
        self.submission_id = 1


# ── .docx export (sheet_docx) ────────────────────────────────────────


def test_render_latex_png_produces_image(tmp_path):
    out = tmp_path / "q.png"
    render_latex_png(LATEX_FRACTION, str(out))
    assert out.exists() and out.stat().st_size > 0


def test_build_docx_embeds_latex_png_and_omits_answer(tmp_path):
    q = SheetQuestionData(
        question_number="1",
        question_type="calculation",
        subject="math",
        question_text=None,
        question_latex=LATEX_FRACTION,
        question_image_path=None,
        source_submission_id=1,
    )
    path = build_sheet_docx([q], "小明", "math", output_dir=str(tmp_path))
    assert os.path.isfile(path)

    doc = Document(path)
    assert len(doc.inline_shapes) >= 1  # LaTeX rendered as PNG
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "错题练习试卷" in text
    assert "第 1 题" in text
    # Answer key must never leak (solution_note is not even a field).
    assert "正确答案" not in text


def test_build_docx_incomplete_falls_back_to_screenshot(tmp_path):
    img_path = tmp_path / "question.jpg"
    Image.new("RGB", (40, 20), color=(255, 255, 255)).save(img_path, "JPEG")

    q = SheetQuestionData(
        question_number="2",
        question_type="word_problem",
        subject="math",
        question_text=None,
        question_latex=None,
        question_image_path=str(img_path),
        source_submission_id=1,
    )
    path = build_sheet_docx([q], "小明", "math", output_dir=str(tmp_path))
    assert os.path.isfile(path)
    doc = Document(path)
    assert len(doc.inline_shapes) >= 1  # screenshot embedded as fallback


# ── generate endpoint: format switch ─────────────────────────────────


@pytest.mark.asyncio
async def test_generate_text_returns_questions_and_docx_url(client, db_session):
    token = await login(client, PHONE)
    child_id = await _get_child_id(client, token)
    await _seed_error_questions(
        db_session,
        child_id,
        "math",
        [
            {"number": "1", "latex": LATEX_FRACTION},
            {"number": "2", "latex": r"x^2 + y^2 = 1"},
            {"number": "3", "latex": None},  # incomplete → image fallback
        ],
    )

    resp = await client.post(
        "/api/error-collections/generate",
        json={"child_id": child_id, "subject": "math", "count": 10, "format": "text"},
        headers=_auth(token),
    )
    assert resp.status_code == 200
    data = resp.json()

    assert data["format"] == "text"
    # count > available → return actual count, no blank padding (AC-X4.5).
    assert data["question_count"] == 3
    assert data["image_url"] is None
    assert data["docx_url"] is not None
    assert len(data["questions"]) == 3

    q = data["questions"][0]
    assert q["subject"] == "math"
    assert q["question_number"]
    assert isinstance(q["source_submission_id"], int)
    # Incomplete question carries a signed fallback image URL.
    incomplete = next(qq for qq in data["questions"] if qq["question_latex"] is None)
    assert incomplete["question_image_path"]


@pytest.mark.asyncio
async def test_generate_text_insufficient_count_returns_actual(client, db_session):
    token = await login(client, PHONE)
    child_id = await _get_child_id(client, token)
    await _seed_error_questions(
        db_session,
        child_id,
        "english",
        [
            {"number": "1", "text": "What is your name?"},
            {"number": "2", "text": "Where are you from?"},
        ],
    )

    resp = await client.post(
        "/api/error-collections/generate",
        json={
            "child_id": child_id,
            "subject": "english",
            "count": 50,
            "format": "text",
        },
        headers=_auth(token),
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["question_count"] == 2
    assert len(data["questions"]) == 2
    assert all(q["question_text"] for q in data["questions"])


@pytest.mark.asyncio
async def test_generate_default_format_is_image_backward_compatible(
    client, db_session, monkeypatch
):
    token = await login(client, PHONE)
    child_id = await _get_child_id(client, token)
    await _seed_error_questions(
        db_session,
        child_id,
        "math",
        [{"number": "1", "latex": LATEX_FRACTION}],
    )

    # Avoid writing a real sheet image to disk during the test.
    monkeypatch.setattr(
        ec_router,
        "compose_sheet",
        lambda errors, child_name, subject: "data/images/sheets/abc.jpg",
    )

    resp = await client.post(
        "/api/error-collections/generate",
        json={"child_id": child_id, "subject": "math", "count": 10},
        headers=_auth(token),
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["format"] == "image"
    assert data["image_url"] is not None
    assert data["question_count"] == 1
    assert data["questions"] is None
    assert data["docx_url"] is None


@pytest.mark.asyncio
async def test_generate_invalid_format_rejected(client, db_session):
    token = await login(client, PHONE)
    child_id = await _get_child_id(client, token)
    resp = await client.post(
        "/api/error-collections/generate",
        json={"child_id": child_id, "subject": "math", "count": 5, "format": "pdf"},
        headers=_auth(token),
    )
    assert resp.status_code == 422


@pytest.mark.asyncio
async def test_generate_text_no_matching_questions(client, db_session):
    token = await login(client, PHONE)
    child_id = await _get_child_id(client, token)
    resp = await client.post(
        "/api/error-collections/generate",
        json={"child_id": child_id, "subject": "english", "count": 5, "format": "text"},
        headers=_auth(token),
    )
    assert resp.status_code == 400


# ── .docx download route ─────────────────────────────────────────────


def _write_docx(filename: str) -> None:
    path = os.path.join(ec_router.SHEET_DOCX_DIR, filename)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    Document().save(path)


def _docx_url(filename: str, expires: int, token: str) -> str:
    return f"/api/sheets/{filename}?token={token}&expires={expires}"


@pytest.mark.asyncio
async def test_serve_docx_valid_signature(client):
    filename = "abc123.docx"
    _write_docx(filename)
    expires = int(time.time()) + 3600
    token = sign("sheets", filename, expires)

    resp = await client.get(_docx_url(filename, expires, token))
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )


@pytest.mark.asyncio
async def test_serve_docx_tampered_token(client):
    filename = "abc124.docx"
    _write_docx(filename)
    expires = int(time.time()) + 3600
    token = sign("sheets", filename, expires)
    tampered = ("a" if not token.endswith("a") else "b") + token[1:]

    resp = await client.get(_docx_url(filename, expires, tampered))
    assert resp.status_code == 403


@pytest.mark.asyncio
async def test_serve_docx_expired_signature(client):
    filename = "abc125.docx"
    _write_docx(filename)
    expires = int(time.time()) - 1
    token = sign("sheets", filename, expires)

    resp = await client.get(_docx_url(filename, expires, token))
    assert resp.status_code == 403


@pytest.mark.asyncio
async def test_serve_docx_path_traversal(client):
    filename = "..%2Fsecret.docx"
    expires = int(time.time()) + 3600
    token = sign("sheets", filename, expires)

    resp = await client.get(_docx_url(filename, expires, token))
    assert resp.status_code == 404


@pytest.mark.asyncio
async def test_serve_docx_missing_file(client):
    filename = "doesnotexist.docx"
    expires = int(time.time()) + 3600
    token = sign("sheets", filename, expires)

    resp = await client.get(_docx_url(filename, expires, token))
    assert resp.status_code == 404
